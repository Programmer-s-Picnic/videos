"""A beginner-friendly document-formatting agent for .docx files."""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


@dataclass
class FormattingPlan:
    font_name: str = "Times New Roman"
    body_size: float = 14
    heading_size: float | None = None
    headings_bold: bool = False
    line_spacing: float | None = None
    margin_inches: float | None = None


def understand_instruction(instruction: str) -> FormattingPlan:
    """Convert a simple English/Hinglish instruction into a formatting plan."""
    text = instruction.lower().strip()
    plan = FormattingPlan()

    known_fonts = {
        "times new roman": "Times New Roman",
        "arial": "Arial",
        "calibri": "Calibri",
        "georgia": "Georgia",
    }
    for key, value in known_fonts.items():
        if key in text:
            plan.font_name = value
            break

    body_match = re.search(r"(?:font\s*)?size\s*(\d+(?:\.\d+)?)", text)
    if body_match:
        plan.body_size = float(body_match.group(1))

    heading_match = re.search(r"headings?[^\d]{0,20}(\d+(?:\.\d+)?)", text)
    if heading_match:
        plan.heading_size = float(heading_match.group(1))

    plan.headings_bold = "heading" in text and "bold" in text

    spacing_match = re.search(r"(?:line\s*)?spacing\s*(\d+(?:\.\d+)?)", text)
    if spacing_match:
        plan.line_spacing = float(spacing_match.group(1))

    if "normal margin" in text or "normal margins" in text:
        plan.margin_inches = 1.0
    else:
        margin_match = re.search(r"margins?\s*(\d+(?:\.\d+)?)", text)
        if margin_match:
            plan.margin_inches = float(margin_match.group(1))

    return plan


def set_run_font(run, name: str, size: float, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def is_heading(paragraph) -> bool:
    return bool(paragraph.style and paragraph.style.name.startswith("Heading"))


def format_paragraph(paragraph, plan: FormattingPlan) -> None:
    heading = is_heading(paragraph)
    size = plan.heading_size if heading and plan.heading_size else plan.body_size
    for run in paragraph.runs:
        set_run_font(
            run,
            plan.font_name,
            size,
            True if heading and plan.headings_bold else None,
        )
    if plan.line_spacing:
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.paragraph_format.line_spacing = plan.line_spacing


def iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested_table in cell.tables:
                yield from iter_table_paragraphs(nested_table)


def iter_all_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        yield from iter_table_paragraphs(table)
    for section in document.sections:
        yield from section.header.paragraphs
        for table in section.header.tables:
            yield from iter_table_paragraphs(table)
        yield from section.footer.paragraphs
        for table in section.footer.tables:
            yield from iter_table_paragraphs(table)


def apply_plan(input_path: Path, output_path: Path, plan: FormattingPlan) -> None:
    document = Document(input_path)
    for paragraph in iter_all_paragraphs(document):
        format_paragraph(paragraph, plan)

    if plan.margin_inches is not None:
        for section in document.sections:
            margin = Inches(plan.margin_inches)
            section.top_margin = margin
            section.right_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Format a Word document using an instruction.")
    parser.add_argument("input", type=Path, help="Input .docx file")
    parser.add_argument("output", type=Path, help="Output .docx file")
    parser.add_argument(
        "--instruction",
        default="Times New Roman font size 14",
        help="Example: Times New Roman size 14, headings 16 bold, line spacing 1.5",
    )
    args = parser.parse_args()

    if args.input.suffix.lower() != ".docx" or args.output.suffix.lower() != ".docx":
        parser.error("Input and output must be .docx files.")
    if not args.input.exists():
        parser.error(f"Input file not found: {args.input}")

    plan = understand_instruction(args.instruction)
    apply_plan(args.input, args.output, plan)
    print(f"Done: {args.output}")
    print(f"Plan: {plan}")


if __name__ == "__main__":
    main()
