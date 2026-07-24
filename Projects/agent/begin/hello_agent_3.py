"""Hello World Agent 3: automate Word document formatting."""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


def observe_document(input_path):
    path = Path(input_path)
    if not path.exists():
        return {"valid": False, "error": "The input file does not exist."}
    if path.suffix.lower() != ".docx":
        return {"valid": False, "error": "Only .docx files are supported."}
    return {"valid": True, "path": path}


def understand_goal(goal):
    goal = goal.lower()
    plan = {"font_name": "Times New Roman", "font_size": 14}
    if "arial" in goal:
        plan["font_name"] = "Arial"
    elif "calibri" in goal:
        plan["font_name"] = "Calibri"
    for size in range(8, 33):
        if str(size) in goal:
            plan["font_size"] = size
            break
    return plan


def set_run_format(run, font_name, font_size):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)


def format_paragraph(paragraph, plan):
    for run in paragraph.runs:
        set_run_format(run, plan["font_name"], plan["font_size"])


def format_table(table, plan):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                format_paragraph(paragraph, plan)
            for nested_table in cell.tables:
                format_table(nested_table, plan)


def format_document(input_path, output_path, plan):
    document = Document(input_path)
    for paragraph in document.paragraphs:
        format_paragraph(paragraph, plan)
    for table in document.tables:
        format_table(table, plan)
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            format_paragraph(paragraph, plan)
        for paragraph in section.footer.paragraphs:
            format_paragraph(paragraph, plan)
    document.save(output_path)
    return output_path


def verify_result(output_path):
    path = Path(output_path)
    return path.exists() and path.stat().st_size > 0


def word_formatting_agent(goal, input_path, output_path):
    print("\nGoal:", goal)
    observation = observe_document(input_path)
    if not observation["valid"]:
        print("Agent stopped:", observation["error"])
        return

    print("Observation: Valid Word document found.")
    plan = understand_goal(goal)
    print("\nPlan:")
    print("Font:", plan["font_name"])
    print("Size:", plan["font_size"])
    print("Output:", output_path)
    print("\nAgent is executing the formatting tools...")
    format_document(observation["path"], output_path, plan)

    if verify_result(output_path):
        print("\nGoal completed successfully.")
        print("Formatted document:", output_path)
    else:
        print("\nThe output could not be verified.")


if __name__ == "__main__":
    instruction = input("Formatting instruction: ")
    word_formatting_agent(instruction, "input.docx", "formatted.docx")
